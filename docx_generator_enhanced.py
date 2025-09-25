#!/usr/bin/env python3
"""
Enhanced Professional DOCX Generator for Resume - Improved Markdown Parsing
Fixed parsing issues and enhanced markdown support
"""

import re
import base64
import io
from typing import Optional, List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class EnhancedDocxGenerator:
    """
    Enhanced DOCX generator with improved markdown parsing and professional styling
    """
    
    def __init__(self):
        self.primary_color = RGBColor(44, 90, 160)  # Professional blue
        self.text_color = RGBColor(51, 51, 51)      # Dark gray
        self.light_color = RGBColor(102, 102, 102)  # Light gray
        
    def create_document(self) -> Document:
        """Create a new document with professional styling"""
        doc = Document()
        
        # Set document margins (0.7 inches)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Create custom styles
        self._create_custom_styles(doc)
        
        return doc
    
    def _create_custom_styles(self, doc: Document):
        """Create custom styles for professional formatting"""
        styles = doc.styles
        
        # Header style for name
        if 'ResumeHeader' not in [s.name for s in styles]:
            header_style = styles.add_style('ResumeHeader', WD_STYLE_TYPE.PARAGRAPH)
            header_font = header_style.font
            header_font.name = 'Segoe UI'
            header_font.size = Pt(20)
            header_font.bold = True
            header_font.color.rgb = self.primary_color
            header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_style.paragraph_format.space_after = Pt(6)
        
        # Contact info style
        if 'ContactInfo' not in [s.name for s in styles]:
            contact_style = styles.add_style('ContactInfo', WD_STYLE_TYPE.PARAGRAPH)
            contact_font = contact_style.font
            contact_font.name = 'Segoe UI'
            contact_font.size = Pt(10)
            contact_font.color.rgb = self.light_color
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_style.paragraph_format.space_after = Pt(12)
        
        # Section header style
        if 'SectionHeader' not in [s.name for s in styles]:
            section_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Segoe UI'
            section_font.size = Pt(12)
            section_font.bold = True
            section_font.color.rgb = self.primary_color
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
            section_style.paragraph_format.keep_with_next = True
        
        # Body text style
        if 'BodyText' not in [s.name for s in styles]:
            body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Segoe UI'
            body_font.size = Pt(10)
            body_font.color.rgb = self.text_color
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Job title style
        if 'JobTitle' not in [s.name for s in styles]:
            job_style = styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
            job_font = job_style.font
            job_font.name = 'Segoe UI'
            job_font.size = Pt(11)
            job_font.bold = True
            job_font.color.rgb = self.text_color
            job_style.paragraph_format.space_after = Pt(3)
        
        # Company/Institution style
        if 'CompanyName' not in [s.name for s in styles]:
            company_style = styles.add_style('CompanyName', WD_STYLE_TYPE.PARAGRAPH)
            company_font = company_style.font
            company_font.name = 'Segoe UI'
            company_font.size = Pt(10)
            company_font.italic = True
            company_font.color.rgb = self.light_color
            company_style.paragraph_format.space_after = Pt(6)
    
    def _add_section_border(self, paragraph):
        """Add a bottom border to section headers"""
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pBdr.set(qn('w:bottom'), 'single')
        pBdr.set(qn('w:sz'), '6')
        pBdr.set(qn('w:space'), '1')
        pBdr.set(qn('w:color'), '2c5aa0')
        pPr.append(pBdr)
    
    def _parse_rich_text(self, text: str, paragraph):
        """Parse text with markdown formatting and add to paragraph with proper styling"""
        # Handle bold, italic, and code formatting while preserving the formatting in Word
        
        # Split by markdown patterns while keeping the delimiters
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
        
        for part in parts:
            if not part:
                continue
                
            run = paragraph.add_run()
            
            # Bold text (**text**)
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run.text = part[2:-2]
                run.bold = True
            # Italic text (*text*)
            elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
                run.text = part[1:-1]
                run.italic = True
            # Code text (`text`)
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                run.text = part[1:-1]
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            # Links [text](url)
            elif re.match(r'\[.*?\]\(.*?\)', part):
                link_match = re.match(r'\[([^\]]+)\]\([^)]+\)', part)
                if link_match:
                    run.text = link_match.group(1)
                    run.font.color.rgb = RGBColor(0, 0, 238)  # Blue color for links
                else:
                    run.text = part
            # Regular text
            else:
                run.text = part
    
    def parse_markdown_to_docx(self, markdown_content: str) -> Document:
        """Enhanced markdown parsing with better formatting support"""
        doc = self.create_document()
        lines = markdown_content.strip().split('\n')
        
        i = 0
        skip_empty_lines = False
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines except when we want to preserve spacing
            if not line:
                if not skip_empty_lines:
                    # Add small spacing
                    pass
                i += 1
                continue
            
            # Main header (name) - must be first line
            if line.startswith('# ') and i < 3:  # Allow some flexibility for position
                name = line[2:].strip()
                header_para = doc.add_paragraph(name, style='ResumeHeader')
                i += 1
                skip_empty_lines = True
                
                # Look for contact info in next few lines
                contact_parts = []
                contact_lines = 0
                while i < len(lines) and contact_lines < 5:
                    next_line = lines[i].strip()
                    if not next_line:
                        i += 1
                        continue
                    if next_line.startswith('#'):
                        break
                    
                    # Check if this looks like contact info
                    contact_indicators = ['@', 'phone', 'linkedin', 'github', '(', ')', '-', '+', 'http', 'www', '.com', '.org']
                    if any(indicator in next_line.lower() for indicator in contact_indicators):
                        # Clean up the line but preserve important formatting
                        clean_line = re.sub(r'^\*+\s*', '', next_line)  # Remove leading asterisks
                        clean_line = re.sub(r'\s*\*+$', '', clean_line)  # Remove trailing asterisks
                        clean_line = re.sub(r'`([^`]+)`', r'\1', clean_line)  # Remove code formatting
                        contact_parts.append(clean_line)
                        contact_lines += 1
                    else:
                        break
                    i += 1
                
                # Add contact info as single line
                if contact_parts:
                    contact_info = ' | '.join(contact_parts)
                    doc.add_paragraph(contact_info, style='ContactInfo')
                
                skip_empty_lines = False
            
            # Section headers (## text)
            elif line.startswith('## '):
                section_name = line[3:].strip()
                section_para = doc.add_paragraph(section_name, style='SectionHeader')
                self._add_section_border(section_para)
                i += 1
                skip_empty_lines = True
            
            # Job/Project titles (bold lines starting with **)
            elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                title = line[2:-2].strip()
                doc.add_paragraph(title, style='JobTitle')
                i += 1
            
            # Company/Institution info (italic lines with single *)
            elif (line.startswith('*') and line.endswith('*') and 
                  not line.startswith('**') and len(line) > 2):
                company = line[1:-1].strip()
                doc.add_paragraph(company, style='CompanyName')
                i += 1
            
            # Bullet points
            elif line.startswith(('- ', '• ', '* ')):
                bullet_text = line[2:].strip()
                para = doc.add_paragraph(style='BodyText')
                
                # Add bullet point
                bullet_run = para.add_run('• ')
                
                # Parse the rest with rich text formatting
                self._parse_rich_text(bullet_text, para)
                i += 1
            
            # Numbered lists
            elif re.match(r'^\d+\.\s+', line):
                list_text = re.sub(r'^\d+\.\s+', '', line)
                para = doc.add_paragraph(style='BodyText')
                
                # Add number
                number_match = re.match(r'^(\d+)\.\s+', line)
                if number_match:
                    num_run = para.add_run(f'{number_match.group(1)}. ')
                
                # Parse the rest with rich text formatting
                self._parse_rich_text(list_text, para)
                i += 1
            
            # Skills and other comma-separated content
            elif (',' in line and 
                  not any(header in line.lower() for header in ['experience', 'education', 'project', 'summary']) and
                  not line.startswith('#')):
                para = doc.add_paragraph(style='BodyText')
                self._parse_rich_text(line, para)
                i += 1
            
            # Regular paragraphs
            else:
                if line.strip():
                    para = doc.add_paragraph(style='BodyText')
                    self._parse_rich_text(line, para)
                i += 1
        
        return doc
    
    def generate_docx_base64(self, markdown_content: str) -> Optional[str]:
        """
        Generate DOCX from markdown and return as base64 string
        
        Args:
            markdown_content: Markdown formatted resume content
            
        Returns:
            Base64 encoded DOCX content or None if generation fails
        """
        try:
            doc = self.parse_markdown_to_docx(markdown_content)
            
            # Save to bytes buffer
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            # Convert to base64
            docx_base64 = base64.b64encode(docx_buffer.read()).decode('utf-8')
            
            print(f"Enhanced DOCX generated successfully, size: {len(docx_base64)} characters")
            return docx_base64
            
        except Exception as e:
            print(f"Error generating enhanced DOCX: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def save_docx_file(self, markdown_content: str, filename: str = "resume.docx") -> bool:
        """
        Generate and save DOCX file to disk
        
        Args:
            markdown_content: Markdown formatted resume content
            filename: Output filename
            
        Returns:
            True if successful, False otherwise
        """
        try:
            doc = self.parse_markdown_to_docx(markdown_content)
            doc.save(filename)
            print(f"Enhanced DOCX saved successfully as {filename}")
            return True
            
        except Exception as e:
            print(f"Error saving enhanced DOCX file: {e}")
            import traceback
            traceback.print_exc()
            return False

# Test the enhanced parser
def test_enhanced_parsing():
    """Test the enhanced parsing with complex markdown"""
    complex_markdown = """# Dr. Emily Rodriguez, Ph.D.
emily.rodriguez@email.com | (555) 123-4567 | LinkedIn: linkedin.com/in/emily-rodriguez
GitHub: github.com/emilyrodriguez | Portfolio: https://emilyrodriguez.dev | San Francisco, CA

## Professional Summary

**Data Scientist** and *Machine Learning Engineer* with **8+ years** of experience in developing `AI solutions` for Fortune 500 companies. Expertise in `Python`, `R`, and **deep learning frameworks**. Published author with [15+ peer-reviewed papers](https://scholar.google.com/citations?user=example).

## Experience

**Senior Data Science Manager**
*Meta (Facebook) | Menlo Park, CA | 2021 - Present*

- Led a team of **12 data scientists** and *ML engineers* in developing recommendation systems
- Improved user engagement metrics by **25%** through advanced `neural networks`
- **Key Achievement**: Deployed production ML models serving `2B+ users` daily
- Technologies: `Python`, `PyTorch`, `TensorFlow`, `SQL`, `Spark`, `Kubernetes`

**Principal Data Scientist**
*Uber Technologies | San Francisco, CA | 2018 - 2021*

- Architected **real-time pricing algorithms** using `machine learning` and *optimization techniques*
- Reduced operational costs by **$50M annually** through predictive analytics
- Published research: [\"Dynamic Pricing in Ride-Sharing Networks\"](https://arxiv.org/abs/example)
- Mentored *15+ junior data scientists* and established **best practices**

**Data Scientist**
*Airbnb | San Francisco, CA | 2016 - 2018*

1. Built fraud detection systems with **99.2% accuracy** using `ensemble methods`
2. Developed A/B testing framework reducing experiment duration by **40%**
3. Created data visualization dashboards used by **500+ employees**

## Education

**Ph.D. in Computer Science**
*Stanford University | Stanford, CA | 2012 - 2016*

- **Dissertation**: "Deep Reinforcement Learning for Autonomous Systems"
- **Advisor**: Prof. Andrew Ng
- **GPA**: 3.95/4.0

**Master of Science in Statistics**
*UC Berkeley | Berkeley, CA | 2010 - 2012*

- *Specialization*: Bayesian Statistics and Machine Learning
- **GPA**: 3.9/4.0

## Technical Skills

**Programming Languages:** `Python`, `R`, `SQL`, `Java`, `Scala`, `JavaScript`
**ML/AI Frameworks:** `PyTorch`, `TensorFlow`, `Scikit-learn`, `XGBoost`, `Keras`
**Big Data Technologies:** `Spark`, `Hadoop`, `Kafka`, `Airflow`, `Databricks`
**Cloud Platforms:** `AWS`, `GCP`, `Azure` - **certified in all three**
**Databases:** `PostgreSQL`, `MongoDB`, `Redis`, `Snowflake`, `BigQuery`

## Publications & Patents

**Recent Publications:**
- [\"Transformer Networks for Time Series Forecasting\"](https://arxiv.org/example) - *NeurIPS 2023*
- [\"Federated Learning in Production Systems\"](https://arxiv.org/example) - **ICML 2022**

**Patents:**
1. **US Patent 11,123,456**: "Method for Real-time Anomaly Detection" (2023)
2. **US Patent 10,987,654**: "Distributed Machine Learning System" (2022)

## Awards & Recognition

- **"Data Scientist of the Year"** - *Data Science Association* (2023)
- **"Outstanding Research Award"** - Stanford University (2016)
- **"Best Paper Award"** - *ICML Conference* (2022)
"""
    
    print("🧪 Testing Enhanced DOCX Parser")
    print("=" * 50)
    
    try:
        generator = EnhancedDocxGenerator()
        
        # Test parsing
        print("📝 Parsing complex markdown...")
        docx_base64 = generator.generate_docx_base64(complex_markdown)
        
        if docx_base64:
            print(f"✅ Success! Generated {len(docx_base64)} characters")
            
            # Save test file
            success = generator.save_docx_file(complex_markdown, "enhanced_test.docx")
            if success:
                print("✅ Test file saved as 'enhanced_test.docx'")
                return True
        else:
            print("❌ Failed to generate DOCX")
            return False
            
    except Exception as e:
        print(f"❌ Test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    test_enhanced_parsing()