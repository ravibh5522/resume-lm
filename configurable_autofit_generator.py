#!/usr/bin/env python3
"""
Configurable Auto-Fit Generator - Adjustable sensitivity controls
Fine-tune auto-fit behavior with customizable parameters
"""

import base64
import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Optional, Tuple, Any

class ConfigurableAutoFitGenerator:
    """
    Auto-fit generator with configurable sensitivity controls
    Allows fine-tuning of when and how much font scaling occurs
    """
    
    def __init__(self, sensitivity_config: Optional[Dict] = None):
        # Color scheme
        self.primary_color = RGBColor(26, 54, 93)
        self.secondary_color = RGBColor(43, 108, 176)
        self.text_color = RGBColor(45, 55, 72)
        self.light_color = RGBColor(74, 85, 104)
        
        # CONFIGURABLE SENSITIVITY PARAMETERS
        # You can adjust these to control auto-fit behavior
        self.sensitivity = sensitivity_config or {
            # CONTENT THRESHOLDS - Adjust these to change when scaling kicks in
            'short_resume_lines': 25,      # Lines below this = no scaling
            'medium_resume_lines': 40,     # Lines below this = light scaling
            'long_resume_lines': 60,       # Lines below this = medium scaling
            'very_long_lines': 80,         # Lines below this = heavy scaling
            
            # SCALING AGGRESSIVENESS - Lower = more gentle, Higher = more aggressive
            'scaling_sensitivity': 1.0,    # Overall scaling multiplier (0.5-2.0)
            'min_font_size': 6,           # Never go below this font size
            'max_font_size': 20,          # Never go above this font size
            
            # CONTENT WEIGHT FACTORS - How much each element affects scaling
            'header_weight': 2.5,         # H2 headers impact
            'bullet_weight': 1.8,         # Bullet points impact  
            'job_entry_weight': 3.0,      # Job entries impact
            'word_weight': 0.12,          # Individual words impact
            'line_weight': 1.2,           # Content lines impact
            
            # FONT SIZE RANGES per scaling level - MORE PROPORTIONAL SIZES!
            'comfortable': {'header': 14, 'section': 11, 'body': 10, 'small': 9},    # Smaller header gap
            'balanced': {'header': 13, 'section': 10, 'body': 9, 'small': 8},        # Better proportions
            'compact': {'header': 12, 'section': 10, 'body': 8, 'small': 7},         # Closer sizing
            'dense': {'header': 11, 'section': 9, 'body': 7, 'small': 6},            # Minimal gap
            'ultra_dense': {'header': 10, 'section': 8, 'body': 6, 'small': 5},      # Very tight
            
            # SPACING PARAMETERS
            'spacing_reduction': 0.8,      # How much to reduce spacing (0.5-1.0)
            'margin_reduction': 0.9,       # How much to reduce margins (0.5-1.0)
        }
        
        print(f"🎛️ Auto-fit sensitivity configured:")
        print(f"   📏 Thresholds: {self.sensitivity['short_resume_lines']}/{self.sensitivity['medium_resume_lines']}/{self.sensitivity['long_resume_lines']}/{self.sensitivity['very_long_lines']} lines")
        print(f"   ⚖️ Sensitivity: {self.sensitivity['scaling_sensitivity']:.1f}x")
        print(f"   📝 Font range: {self.sensitivity['min_font_size']}-{self.sensitivity['max_font_size']}pt")
    
    def analyze_content_with_sensitivity(self, markdown_text: str) -> Dict[str, Any]:
        """
        Analyze content using configurable sensitivity parameters
        """
        analysis = {
            'total_chars': len(markdown_text),
            'content_lines': len([line for line in markdown_text.split('\n') if line.strip()]),
            'h2_headers': len(re.findall(r'^##\s', markdown_text, re.MULTILINE)),
            'bullet_points': len(re.findall(r'^[\s]*[-•*]\s+', markdown_text, re.MULTILINE)),
            'job_entries': len(re.findall(r'^\*\*[^*]+\*\*\s*$', markdown_text, re.MULTILINE)),
            'estimated_words': len(markdown_text.split()),
        }
        
        # Calculate density using configurable weights
        density_score = (
            analysis['content_lines'] * self.sensitivity['line_weight'] +
            analysis['h2_headers'] * self.sensitivity['header_weight'] +
            analysis['bullet_points'] * self.sensitivity['bullet_weight'] +
            analysis['job_entries'] * self.sensitivity['job_entry_weight'] +
            analysis['estimated_words'] * self.sensitivity['word_weight']
        )
        
        # Apply overall sensitivity multiplier
        density_score *= self.sensitivity['scaling_sensitivity']
        
        analysis['density_score'] = int(density_score)
        analysis['sensitivity_applied'] = self.sensitivity['scaling_sensitivity']
        
        print(f"🔍 Sensitivity Analysis:")
        print(f"   📊 Lines: {analysis['content_lines']}, Words: {analysis['estimated_words']}")
        print(f"   🎯 Weighted density: {density_score:.1f} (sensitivity: {self.sensitivity['scaling_sensitivity']:.1f}x)")
        
        return analysis
    
    def calculate_configurable_scaling(self, analysis: Dict[str, Any]) -> Dict[str, Any]:
        """
        Calculate scaling using configurable thresholds and parameters
        """
        content_lines = analysis['content_lines']
        
        # Use configurable thresholds
        if content_lines <= self.sensitivity['short_resume_lines']:
            level = 'comfortable'
            factor = 1.0
            class_name = 'COMFORTABLE'
            spacing_after = 6
            line_spacing = 1.15
            margins = 0.8
            
        elif content_lines <= self.sensitivity['medium_resume_lines']:
            level = 'balanced'
            factor = 0.9
            class_name = 'BALANCED'
            spacing_after = 4
            line_spacing = 1.1
            margins = 0.7
            
        elif content_lines <= self.sensitivity['long_resume_lines']:
            level = 'compact'
            factor = 0.8
            class_name = 'COMPACT'
            spacing_after = 3
            line_spacing = 1.05
            margins = 0.6
            
        elif content_lines <= self.sensitivity['very_long_lines']:
            level = 'dense'
            factor = 0.7
            class_name = 'DENSE'
            spacing_after = 2
            line_spacing = 1.0
            margins = 0.5
            
        else:
            level = 'ultra_dense'
            factor = 0.6
            class_name = 'ULTRA_DENSE'
            spacing_after = 1
            line_spacing = 0.95
            margins = 0.4
        
        # Apply sensitivity adjustments to spacing and margins
        spacing_after = max(1, int(spacing_after * self.sensitivity['spacing_reduction']))
        margins *= self.sensitivity['margin_reduction']
        
        # Get font sizes for this level
        fonts = self.sensitivity[level].copy()
        
        # Ensure fonts stay within min/max bounds
        for font_type in fonts:
            fonts[font_type] = max(
                self.sensitivity['min_font_size'],
                min(self.sensitivity['max_font_size'], fonts[font_type])
            )
        
        scale = {
            'factor': factor,
            'class': class_name,
            'level': level,
            'header': fonts['header'],
            'section': fonts['section'], 
            'body': fonts['body'],
            'small': fonts['small'],
            'line_spacing': line_spacing,
            'spacing_after': spacing_after,
            'margins': margins
        }
        
        print(f"🎯 Configurable scaling: {class_name} - Body: {fonts['body']}pt, Spacing: {spacing_after}pt")
        
        return scale
    
    def generate_configurable_docx_base64(self, markdown_text: str) -> str:
        """
        Generate DOCX with configurable auto-fit scaling
        """
        try:
            print("🎛️ Starting configurable auto-fit DOCX generation...")
            
            # Analyze content with sensitivity
            analysis = self.analyze_content_with_sensitivity(markdown_text)
            
            # Calculate scaling with configurable parameters
            scaling = self.calculate_configurable_scaling(analysis)
            
            # Clean markdown
            clean_markdown = self._clean_markdown(markdown_text)
            
            # Create document with calculated scaling
            document = Document()
            
            # Set margins based on scaling
            for section in document.sections:
                margin_inches = scaling['margins']
                section.top_margin = Inches(margin_inches)
                section.bottom_margin = Inches(margin_inches)
                section.left_margin = Inches(margin_inches)
                section.right_margin = Inches(margin_inches)
            
            # Setup styles with calculated font sizes
            self._setup_configurable_styles(document, scaling)
            
            # Process content
            self._process_content_with_scaling(document, clean_markdown, scaling)
            
            # Generate and encode
            doc_buffer = io.BytesIO()
            document.save(doc_buffer)
            doc_buffer.seek(0)
            
            docx_base64 = base64.b64encode(doc_buffer.read()).decode('utf-8')
            print(f"✅ Configurable auto-fit DOCX: {len(docx_base64)} chars, Level: {scaling['class']}")
            
            return docx_base64
            
        except Exception as e:
            print(f"❌ Configurable auto-fit DOCX failed: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    def _clean_markdown(self, markdown_text: str) -> str:
        """Clean markdown for processing"""
        text = markdown_text.strip()
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+$', '', text, flags=re.MULTILINE)
        return text
    
    def _setup_configurable_styles(self, document, scaling: Dict[str, Any]) -> None:
        """Setup document styles using configurable scaling"""
        styles = document.styles
        
        # Normal style
        normal = styles['Normal']
        normal.font.name = 'Segoe UI'
        normal.font.size = Pt(scaling['body'])
        normal.font.color.rgb = self.text_color
        normal.paragraph_format.space_after = Pt(scaling['spacing_after'])
        normal.paragraph_format.line_spacing = scaling['line_spacing']
        
        # Header style
        if 'ConfigurableHeader' not in [s.name for s in styles]:
            header_style = styles.add_style('ConfigurableHeader', WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.name = 'Segoe UI'
            header_style.font.size = Pt(scaling['header'])
            header_style.font.bold = True
            header_style.font.color.rgb = self.primary_color
            header_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_style.paragraph_format.space_after = Pt(scaling['spacing_after'])
        
        # Section style
        if 'ConfigurableSection' not in [s.name for s in styles]:
            section_style = styles.add_style('ConfigurableSection', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Segoe UI'
            section_style.font.size = Pt(scaling['section'])
            section_style.font.bold = True
            section_style.font.color.rgb = self.secondary_color
            section_style.paragraph_format.space_after = Pt(scaling['spacing_after'])
            section_style.paragraph_format.space_before = Pt(scaling['spacing_after'] * 1.5)
    
    def _process_content_with_scaling(self, document, markdown: str, scaling: Dict[str, Any]) -> None:
        """Process content with scaling applied"""
        lines = markdown.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                # Main header
                header_text = line[2:].strip()
                p = document.add_paragraph(header_text, style='ConfigurableHeader')
                
            elif line.startswith('## '):
                # Section header
                section_text = line[3:].strip()
                p = document.add_paragraph(section_text, style='ConfigurableSection')
                
            elif line.startswith(('- ', '• ', '* ')):
                # Bullet point
                bullet_text = re.sub(r'^[-•*]\s+', '', line)
                p = document.add_paragraph()
                p.paragraph_format.space_after = Pt(scaling['spacing_after'] // 2)
                p.add_run('• ')
                self._add_rich_text(bullet_text, p)
                
            else:
                # Regular paragraph
                p = document.add_paragraph()
                p.paragraph_format.space_after = Pt(scaling['spacing_after'])
                self._add_rich_text(line, p)
    
    def _add_rich_text(self, text: str, paragraph) -> None:
        """Add rich text with formatting"""
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        
        for part in parts:
            if not part:
                continue
                
            run = paragraph.add_run()
            
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run.text = part[1:-1]
                run.italic = True
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                run.text = part[1:-1]
                run.font.name = 'Consolas'
            else:
                run.text = part

# PREDEFINED SENSITIVITY CONFIGURATIONS
# You can use these or create your own

SENSITIVITY_PROFILES = {
    'conservative': {
        # Very gentle scaling - keeps fonts larger
        'short_resume_lines': 30,
        'medium_resume_lines': 50, 
        'long_resume_lines': 70,
        'very_long_lines': 90,
        'scaling_sensitivity': 0.7,
        'min_font_size': 7,
        'max_font_size': 20,
        'spacing_reduction': 0.9,
        'margin_reduction': 0.95,
        'comfortable': {'header': 20, 'section': 14, 'body': 11, 'small': 10},
        'balanced': {'header': 18, 'section': 12, 'body': 10, 'small': 9},
        'compact': {'header': 16, 'section': 11, 'body': 9, 'small': 8},
        'dense': {'header': 15, 'section': 10, 'body': 8, 'small': 7},
        'ultra_dense': {'header': 14, 'section': 9, 'body': 7, 'small': 6},
        'header_weight': 2.0, 'bullet_weight': 1.5, 'job_entry_weight': 2.5, 'word_weight': 0.10, 'line_weight': 1.0,
    },
    
    'aggressive': {
        # More aggressive scaling - shrinks fonts quickly  
        'short_resume_lines': 20,
        'medium_resume_lines': 30,
        'long_resume_lines': 45, 
        'very_long_lines': 60,
        'scaling_sensitivity': 1.3,
        'min_font_size': 5,
        'max_font_size': 18,
        'spacing_reduction': 0.6,
        'margin_reduction': 0.7,
        'comfortable': {'header': 16, 'section': 11, 'body': 9, 'small': 8},
        'balanced': {'header': 14, 'section': 10, 'body': 8, 'small': 7},
        'compact': {'header': 13, 'section': 9, 'body': 7, 'small': 6},
        'dense': {'header': 12, 'section': 8, 'body': 6, 'small': 5},
        'ultra_dense': {'header': 11, 'section': 7, 'body': 5, 'small': 5},
        'header_weight': 3.0, 'bullet_weight': 2.2, 'job_entry_weight': 3.5, 'word_weight': 0.15, 'line_weight': 1.4,
    },
    
    'balanced': {
        # Default balanced scaling
        'short_resume_lines': 25,
        'medium_resume_lines': 40,
        'long_resume_lines': 60,
        'very_long_lines': 80,
        'scaling_sensitivity': 1.0,
        'min_font_size': 6,
        'max_font_size': 20,
        'spacing_reduction': 0.8,
        'margin_reduction': 0.9,
        'comfortable': {'header': 18, 'section': 12, 'body': 10, 'small': 9},
        'balanced': {'header': 16, 'section': 11, 'body': 9, 'small': 8},
        'compact': {'header': 15, 'section': 10, 'body': 8, 'small': 7},
        'dense': {'header': 14, 'section': 9, 'body': 7, 'small': 6},
        'ultra_dense': {'header': 13, 'section': 8, 'body': 6, 'small': 5},
        'header_weight': 2.5, 'bullet_weight': 1.8, 'job_entry_weight': 3.0, 'word_weight': 0.12, 'line_weight': 1.2,
    }
}

# Test function with different sensitivity levels
def test_sensitivity_levels():
    """Test different sensitivity configurations"""
    
    test_resume = """# Sarah Johnson
sarah.johnson@email.com | (555) 987-6543

## Professional Summary
**Product Manager** with **6+ years** of experience leading *cross-functional teams* to deliver **innovative software products**. Expertise in `agile methodologies`.

## Experience
**Senior Product Manager**
*HubSpot | Cambridge, MA | 2021 - Present*
- Led product development for **CRM platform** serving `100K+ businesses`
- **Key Achievement**: Increased user engagement by **40%** through *UX optimization*
- Managed **product roadmap** and coordinated with *engineering teams*

**Product Manager**
*Shopify | Boston, MA | 2018 - 2021*
- Launched **e-commerce analytics** dashboard used by `50K+ merchants`
- **Impact**: Improved merchant retention by **25%**

## Education
**MBA in Technology Management**
*MIT Sloan School of Management | Cambridge, MA | 2016 - 2018*

## Skills
**Product Management:** `Roadmapping`, **Agile/Scrum**, *User Research*
**Technical:** **SQL**, `Python`, *JavaScript*, **Figma**"""

    print("🧪 Testing Different Sensitivity Levels")
    print("=" * 60)
    
    for profile_name, config in SENSITIVITY_PROFILES.items():
        print(f"\n🎛️ Testing {profile_name.upper()} sensitivity:")
        print("-" * 30)
        
        try:
            generator = ConfigurableAutoFitGenerator(config)
            result = generator.generate_configurable_docx_base64(test_resume)
            print(f"✅ {profile_name.capitalize()}: {len(result):,} characters")
            
        except Exception as e:
            print(f"❌ {profile_name.capitalize()} failed: {e}")
    
    print("\n" + "=" * 60)
    print("🎉 Sensitivity testing complete!")
    print("💡 Choose the profile that best fits your needs:")
    print("   🐌 Conservative: Keeps fonts larger, gentle scaling")
    print("   ⚖️ Balanced: Good middle ground (default)")
    print("   🚀 Aggressive: Shrinks fonts quickly, fits more content")

if __name__ == "__main__":
    test_sensitivity_levels()