#!/usr/bin/env python3
"""
Professional DOCX Generator for Resume - Microsoft Word Format
Generates high-quality, ATS-optimized Word documents with modern styling
"""

import re
import base64
import io
from typing import Optional, List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class ProfessionalDocxGenerator:
    """
    Professional DOCX generator with modern styling and ATS optimization
    """
    
    def __init__(self):
        self.primary_color = RGBColor(44, 90, 160)  # Professional blue
        self.text_color = RGBColor(51, 51, 51)      # Dark gray
        self.light_color = RGBColor(102, 102, 102)  # Light gray
        
    def create_document(self) -> Document:
        """Create a new document with professional styling"""
        doc = Document()
        
        # Set document margins (0.7 inches)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Create custom styles
        self._create_custom_styles(doc)
        
        return doc
    
    def _create_custom_styles(self, doc: Document):
        """Create custom styles for professional formatting"""
        styles = doc.styles
        
        # Header style for name
        if 'ResumeHeader' not in [s.name for s in styles]:
            header_style = styles.add_style('ResumeHeader', WD_STYLE_TYPE.PARAGRAPH)
            header_font = header_style.font
            header_font.name = 'Segoe UI'
            header_font.size = Pt(20)
            header_font.bold = True
            header_font.color.rgb = self.primary_color
            header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_style.paragraph_format.space_after = Pt(6)
        
        # Contact info style
        if 'ContactInfo' not in [s.name for s in styles]:
            contact_style = styles.add_style('ContactInfo', WD_STYLE_TYPE.PARAGRAPH)
            contact_font = contact_style.font
            contact_font.name = 'Segoe UI'
            contact_font.size = Pt(10)
            contact_font.color.rgb = self.light_color
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_style.paragraph_format.space_after = Pt(12)
        
        # Section header style
        if 'SectionHeader' not in [s.name for s in styles]:
            section_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Segoe UI'
            section_font.size = Pt(12)
            section_font.bold = True
            section_font.color.rgb = self.primary_color
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
            # Add bottom border
            section_style.paragraph_format.keep_with_next = True
        
        # Body text style
        if 'BodyText' not in [s.name for s in styles]:
            body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Segoe UI'
            body_font.size = Pt(10)
            body_font.color.rgb = self.text_color
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Job title style
        if 'JobTitle' not in [s.name for s in styles]:
            job_style = styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
            job_font = job_style.font
            job_font.name = 'Segoe UI'
            job_font.size = Pt(11)
            job_font.bold = True
            job_font.color.rgb = self.text_color
            job_style.paragraph_format.space_after = Pt(3)
        
        # Company/Institution style
        if 'CompanyName' not in [s.name for s in styles]:
            company_style = styles.add_style('CompanyName', WD_STYLE_TYPE.PARAGRAPH)
            company_font = company_style.font
            company_font.name = 'Segoe UI'
            company_font.size = Pt(10)
            company_font.italic = True
            company_font.color.rgb = self.light_color
            company_style.paragraph_format.space_after = Pt(6)
    
    def _add_section_border(self, paragraph):
        """Add a bottom border to section headers"""
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pBdr.set(qn('w:bottom'), 'single')
        pBdr.set(qn('w:sz'), '6')
        pBdr.set(qn('w:space'), '1')
        pBdr.set(qn('w:color'), '2c5aa0')
        pPr.append(pBdr)
    
    def parse_markdown_to_docx(self, markdown_content: str) -> Document:
        """Convert markdown resume to professional DOCX format"""
        doc = self.create_document()
        lines = markdown_content.strip().split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Main header (name)
            if line.startswith('# ') and i == 0:
                name = line[2:].strip()
                header_para = doc.add_paragraph(name, style='ResumeHeader')
                i += 1
                
                # Look for contact info in next few lines
                contact_parts = []
                while i < len(lines) and i < 5:  # Check next 5 lines for contact info
                    next_line = lines[i].strip()
                    if not next_line or next_line.startswith('#'):
                        break
                    # Remove markdown formatting and collect contact info
                    clean_line = re.sub(r'[*_`]', '', next_line)
                    if any(indicator in clean_line.lower() for indicator in ['@', 'phone', 'linkedin', 'github', '(', ')', '-']):
                        contact_parts.append(clean_line)
                    i += 1
                
                # Add contact info as single line
                if contact_parts:
                    contact_info = ' | '.join(contact_parts)
                    doc.add_paragraph(contact_info, style='ContactInfo')
            
            # Section headers
            elif line.startswith('## '):
                section_name = line[3:].strip()
                section_para = doc.add_paragraph(section_name, style='SectionHeader')
                self._add_section_border(section_para)
                i += 1
            
            # Job/Project titles (bold lines)
            elif line.startswith('**') and line.endswith('**'):
                title = line[2:-2].strip()
                doc.add_paragraph(title, style='JobTitle')
                i += 1
            
            # Company/Institution info (italic lines)
            elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
                company = line[1:-1].strip()
                doc.add_paragraph(company, style='CompanyName')
                i += 1
            
            # Bullet points
            elif line.startswith('- ') or line.startswith('• '):
                bullet_text = line[2:].strip()
                # Remove markdown formatting
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', bullet_text)  # Bold
                clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)      # Italic
                clean_text = re.sub(r'`(.*?)`', r'\1', clean_text)        # Code
                
                para = doc.add_paragraph(style='BodyText')
                run = para.runs[0] if para.runs else para.add_run()
                run.text = f"• {clean_text}"
                i += 1
            
            # Skills (comma-separated)
            elif ',' in line and not any(x in line.lower() for x in ['experience', 'education', 'project']):
                # This might be a skills line
                skills_text = re.sub(r'[*_`]', '', line)  # Remove markdown
                doc.add_paragraph(skills_text, style='BodyText')
                i += 1
            
            # Regular text
            else:
                # Clean markdown formatting
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Bold
                clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)  # Italic
                clean_text = re.sub(r'`(.*?)`', r'\1', clean_text)    # Code
                clean_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean_text)  # Links
                
                if clean_text.strip():
                    doc.add_paragraph(clean_text, style='BodyText')
                i += 1
        
        return doc
    
    def generate_docx_base64(self, markdown_content: str) -> Optional[str]:
        """
        Generate DOCX from markdown and return as base64 string
        
        Args:
            markdown_content: Markdown formatted resume content
            
        Returns:
            Base64 encoded DOCX content or None if generation fails
        """
        try:
            doc = self.parse_markdown_to_docx(markdown_content)
            
            # Save to bytes buffer
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            # Convert to base64
            docx_base64 = base64.b64encode(docx_buffer.read()).decode('utf-8')
            
            print(f"DOCX generated successfully, size: {len(docx_base64)} characters")
            return docx_base64
            
        except Exception as e:
            print(f"Error generating DOCX: {e}")
            return None
    
    def save_docx_file(self, markdown_content: str, filename: str = "resume.docx") -> bool:
        """
        Generate and save DOCX file to disk
        
        Args:
            markdown_content: Markdown formatted resume content
            filename: Output filename
            
        Returns:
            True if successful, False otherwise
        """
        try:
            doc = self.parse_markdown_to_docx(markdown_content)
            doc.save(filename)
            print(f"DOCX saved successfully as {filename}")
            return True
            
        except Exception as e:
            print(f"Error saving DOCX file: {e}")
            return False

# Test function
def test_docx_generation():
    """Test the DOCX generation with sample data"""
    sample_markdown = """# John Doe
john.doe@email.com | (555) 123-4567 | LinkedIn: linkedin.com/in/johndoe | San Francisco, CA

## Professional Summary

Experienced software engineer with 5+ years in full-stack development, specializing in Python, React, and cloud technologies. Proven track record of delivering scalable solutions and leading cross-functional teams.

## Experience

**Senior Software Engineer**
*Tech Innovations Inc. | San Francisco, CA | 2020 - Present*

- Led development of microservices architecture serving 1M+ users daily
- Improved system performance by 40% through database optimization
- Mentored junior developers and conducted code reviews
- Technologies: Python, Django, React, PostgreSQL, AWS

**Software Engineer**
*StartupCorp | San Francisco, CA | 2018 - 2020*

- Developed RESTful APIs and responsive web applications
- Collaborated with designers to implement pixel-perfect UI components
- Reduced deployment time by 60% through CI/CD automation
- Technologies: JavaScript, Node.js, MongoDB, Docker

## Education

**Bachelor of Science in Computer Science**
*University of California, Berkeley | 2014 - 2018*

- GPA: 3.8/4.0
- Relevant Coursework: Data Structures, Algorithms, Software Engineering

## Skills

**Technical:** Python, JavaScript, React, Django, Node.js, PostgreSQL, MongoDB, AWS, Docker, Kubernetes
**Tools:** Git, Jenkins, Jira, Figma
**Soft Skills:** Leadership, Communication, Problem-solving, Agile Development

## Projects

**E-commerce Platform**
- Built full-stack e-commerce solution with payment integration
- Technologies: React, Django, Stripe API, PostgreSQL
- GitHub: github.com/johndoe/ecommerce-platform
"""

    generator = ProfessionalDocxGenerator()
    
    # Test base64 generation
    docx_base64 = generator.generate_docx_base64(sample_markdown)
    if docx_base64:
        print(f"✅ Base64 generation successful, length: {len(docx_base64)}")
    else:
        print("❌ Base64 generation failed")
    
    # Test file saving
    success = generator.save_docx_file(sample_markdown, "test_resume.docx")
    if success:
        print("✅ File saving successful")
    else:
        print("❌ File saving failed")

if __name__ == "__main__":
    test_docx_generation()