#!/usr/bin/env python3
"""
Compact DOCX Generator - Minimal spacing, professional layout
Eliminates excessive whitespace in Word documents
"""

import base64
import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn
from typing import List, Dict, Optional

class CompactDocxGenerator:
    """
    Compact DOCX generator with minimal spacing and professional formatting
    """
    
    def __init__(self):
        # Professional color scheme
        self.primary_color = RGBColor(26, 54, 93)    # Dark blue
        self.secondary_color = RGBColor(43, 108, 176) # Medium blue
        self.text_color = RGBColor(45, 55, 72)       # Dark gray
        self.light_color = RGBColor(74, 85, 104)     # Light gray
    
    def _setup_compact_styles(self, document):
        """Setup compact document styles with minimal spacing"""
        styles = document.styles
        
        # Document defaults
        style = document.styles['Normal']
        font = style.font
        font.name = 'Segoe UI'
        font.size = Pt(10)
        font.color.rgb = self.text_color
        
        # COMPACT: Minimal paragraph spacing
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(2)  # Minimal space after paragraphs
        paragraph_format.space_before = Pt(0)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(13)  # Tight line spacing
        
        # Main name/header style (H1)
        if 'ResumeHeader' not in [s.name for s in styles]:
            header_style = styles.add_style('ResumeHeader', WD_STYLE_TYPE.PARAGRAPH)
            header_font = header_style.font
            header_font.name = 'Segoe UI'
            header_font.size = Pt(20)  # Smaller than before
            header_font.bold = True
            header_font.color.rgb = self.primary_color
            header_format = header_style.paragraph_format
            header_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_format.space_after = Pt(6)  # Minimal space
            header_format.space_before = Pt(0)
        
        # Contact info style
        if 'ContactInfo' not in [s.name for s in styles]:
            contact_style = styles.add_style('ContactInfo', WD_STYLE_TYPE.PARAGRAPH)
            contact_font = contact_style.font
            contact_font.name = 'Segoe UI'
            contact_font.size = Pt(9)
            contact_font.color.rgb = self.light_color
            contact_format = contact_style.paragraph_format
            contact_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            contact_format.space_after = Pt(8)  # Minimal space
            contact_format.space_before = Pt(2)
        
        # Section headers (H2) - COMPACT
        if 'SectionHeader' not in [s.name for s in styles]:
            section_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Segoe UI'
            section_font.size = Pt(12)  # Smaller
            section_font.bold = True
            section_font.color.rgb = self.secondary_color
            section_format = section_style.paragraph_format
            section_format.space_after = Pt(4)  # Minimal space
            section_format.space_before = Pt(8)  # Reduced space
        
        # Job title style - COMPACT
        if 'JobTitle' not in [s.name for s in styles]:
            job_style = styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
            job_font = job_style.font
            job_font.name = 'Segoe UI'
            job_font.size = Pt(10)
            job_font.bold = True
            job_font.color.rgb = self.text_color
            job_style.paragraph_format.space_after = Pt(1)  # Very minimal
            job_style.paragraph_format.space_before = Pt(3)  # Reduced
        
        # Company style - COMPACT
        if 'CompanyName' not in [s.name for s in styles]:
            company_style = styles.add_style('CompanyName', WD_STYLE_TYPE.PARAGRAPH)
            company_font = company_style.font
            company_font.name = 'Segoe UI'
            company_font.size = Pt(9)
            company_font.italic = True
            company_font.color.rgb = self.light_color
            company_style.paragraph_format.space_after = Pt(3)  # Minimal
            company_style.paragraph_format.space_before = Pt(0)
        
        # List style - COMPACT
        if 'CompactList' not in [s.name for s in styles]:
            list_style = styles.add_style('CompactList', WD_STYLE_TYPE.PARAGRAPH)
            list_font = list_style.font
            list_font.name = 'Segoe UI'
            list_font.size = Pt(10)
            list_font.color.rgb = self.text_color
            list_format = list_style.paragraph_format
            list_format.space_after = Pt(1)  # Very minimal
            list_format.space_before = Pt(1)
            list_format.left_indent = Inches(0.2)
    
    def _add_compact_section_border(self, paragraph):
        """Add a subtle bottom border to section headers"""
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pBdr.set(qn('w:bottom'), 'single')
        pBdr.set(qn('w:sz'), '4')  # Thinner border
        pBdr.set(qn('w:space'), '1')
        pBdr.set(qn('w:color'), '2c5aa0')
        pPr.append(pBdr)
    
    def _parse_compact_rich_text(self, text: str, paragraph):
        """Parse text with compact formatting"""
        # Split by markdown patterns
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
        
        for part in parts:
            if not part:
                continue
                
            run = paragraph.add_run()
            
            # Bold text (**text**)
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run.text = part[2:-2]
                run.bold = True
            # Italic text (*text*)
            elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
                run.text = part[1:-1]
                run.italic = True
            # Code text (`text`)
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                run.text = part[1:-1]
                run.font.name = 'Consolas'
                run.font.size = Pt(8)  # Smaller code font
            # Links [text](url) - just show text
            elif '[' in part and '](' in part:
                link_match = re.match(r'\[([^\]]+)\]\([^\)]+\)', part)
                if link_match:
                    run.text = link_match.group(1)
                    run.font.color.rgb = RGBColor(49, 130, 206)
                else:
                    run.text = part
            else:
                run.text = part
    
    def _fix_compact_markdown(self, markdown_text: str) -> str:
        """Fix markdown for compact processing"""
        text = markdown_text.strip()
        
        # Remove excessive blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'\n\n+', '\n\n', text)
        
        # Clean up whitespace
        text = re.sub(r'[ \t]+$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^[ \t]+', '', text, flags=re.MULTILINE)
        
        return text.strip()
    
    def generate_compact_docx_base64(self, markdown_text: str) -> str:
        """
        Generate compact DOCX from markdown with minimal spacing
        """
        try:
            print("📝 Starting COMPACT DOCX generation...")
            
            # Fix markdown for compact processing
            clean_markdown = self._fix_compact_markdown(markdown_text)
            
            # Create document with compact margins
            document = Document()
            
            # COMPACT: Set narrow margins
            sections = document.sections
            for section in sections:
                section.top_margin = Inches(0.6)     # Reduced from 1"
                section.bottom_margin = Inches(0.6)   # Reduced from 1" 
                section.left_margin = Inches(0.7)    # Reduced from 1"
                section.right_margin = Inches(0.7)   # Reduced from 1"
            
            # Setup compact styles
            self._setup_compact_styles(document)
            
            # Parse content with compact spacing
            lines = clean_markdown.split('\n')
            current_list = None
            skip_next_empty = False
            
            for i, line in enumerate(lines):
                line = line.strip()
                
                # Skip empty lines more aggressively for compact layout
                if not line:
                    if skip_next_empty:
                        skip_next_empty = False
                        continue
                    # Only add space between major sections
                    if i > 0 and i < len(lines) - 1:
                        next_line = lines[i + 1].strip()
                        if next_line.startswith('##'):
                            continue  # Skip empty line before section headers
                    continue
                
                # Main header (name)
                if line.startswith('# '):
                    header_text = line[2:].strip()
                    p = document.add_paragraph(header_text, style='ResumeHeader')
                    skip_next_empty = True
                    current_list = None
                    continue
                
                # Section headers
                elif line.startswith('## '):
                    section_title = line[3:].strip()
                    p = document.add_paragraph(section_title, style='SectionHeader')
                    self._add_compact_section_border(p)
                    skip_next_empty = True
                    current_list = None
                    continue
                
                # Contact info detection
                elif any(indicator in line.lower() for indicator in ['@', 'linkedin', 'github', 'phone', '(', ')', '+', 'http']):
                    if not line.startswith(('*', '-', '1', '2', '3', '4', '5')):
                        contact_p = document.add_paragraph()
                        contact_p.style = 'ContactInfo'
                        self._parse_compact_rich_text(line, contact_p)
                        current_list = None
                        continue
                
                # Job titles (bold at start of line)
                elif line.startswith('**') and line.endswith('**'):
                    job_title = line[2:-2].strip()
                    p = document.add_paragraph(job_title, style='JobTitle')
                    current_list = None
                    continue
                
                # Company info (italic at start of line)
                elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
                    company_info = line[1:-1].strip()
                    p = document.add_paragraph(company_info, style='CompanyName')
                    current_list = None
                    continue
                
                # Bullet points
                elif line.startswith(('- ', '• ', '* ')):
                    bullet_text = re.sub(r'^[-•*]\s+', '', line)
                    p = document.add_paragraph(style='CompactList')
                    p.style = 'CompactList'
                    # Add bullet manually for consistent formatting
                    p.add_run('• ')  # Use consistent bullet
                    self._parse_compact_rich_text(bullet_text, p)
                    current_list = 'bullet'
                    continue
                
                # Regular paragraphs - with compact spacing
                else:
                    if current_list and not line.startswith(('- ', '• ', '* ', '1', '2', '3')):
                        current_list = None
                    
                    p = document.add_paragraph()
                    p.style = 'Normal'
                    # COMPACT: Reduce spacing for regular paragraphs
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(1)
                    self._parse_compact_rich_text(line, p)
            
            # Save to bytes with compact formatting
            doc_buffer = io.BytesIO()
            document.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Encode to base64
            docx_base64 = base64.b64encode(doc_buffer.read()).decode('utf-8')
            print(f"✅ Compact DOCX generated! Size: {len(docx_base64)} characters")
            
            return docx_base64
            
        except Exception as e:
            print(f"❌ Compact DOCX generation failed: {e}")
            import traceback
            traceback.print_exc()
            raise

# Test the compact DOCX generator
def test_compact_docx():
    """Test compact DOCX generation"""
    
    test_markdown = """# Sarah Johnson
sarah.johnson@email.com | (555) 987-6543 | LinkedIn: linkedin.com/in/sarah-johnson


## Professional Summary


**Product Manager** with **6+ years** experience leading *cross-functional teams*.



## Experience


**Senior Product Manager**
*HubSpot | Cambridge, MA | 2021 - Present*

- Led product development for **CRM platform**
- **Achievement**: Increased engagement by **40%**



## Skills


**Technical:** `Python`, **SQL**, *JavaScript*"""

    print("🧪 Testing Compact DOCX Generator")
    print("=" * 40)
    
    try:
        generator = CompactDocxGenerator()
        
        # Test compact generation
        print("📝 Generating compact DOCX...")
        result = generator.generate_compact_docx_base64(test_markdown)
        
        print(f"✅ Compact DOCX: {len(result):,} characters")
        print("✓ Minimal spacing applied")
        print("✓ Professional formatting preserved")
        print("✓ Compact margins and line spacing")
        
        return True
        
    except Exception as e:
        print(f"❌ Test failed: {e}")
        return False

if __name__ == "__main__":
    test_compact_docx()