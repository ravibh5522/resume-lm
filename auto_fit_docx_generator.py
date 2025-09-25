#!/usr/bin/env python3
"""
Auto-Fit DOCX Generator - Dynamically adjusts font size to fit single page
Analyzes content length and automatically scales DOCX fonts accordingly
"""

import base64
import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Optional, Tuple, Any

class AutoFitDocxGenerator:
    """
    Auto-fit DOCX generator that dynamically adjusts font sizes for single-page layout
    """
    
    def __init__(self):
        self.primary_color = RGBColor(26, 54, 93)
        self.secondary_color = RGBColor(43, 108, 176)
        self.text_color = RGBColor(45, 55, 72)
        self.light_color = RGBColor(74, 85, 104)
    
    def analyze_docx_content(self, markdown_text: str) -> Dict[str, int]:
        """
        Analyze content for DOCX auto-fitting
        """
        analysis = {
            'total_chars': len(markdown_text),
            'content_lines': len([line for line in markdown_text.split('\n') if line.strip()]),
            'h2_headers': len(re.findall(r'^##\s', markdown_text, re.MULTILINE)),
            'bullet_points': len(re.findall(r'^[\s]*[-•*]\s+', markdown_text, re.MULTILINE)),
            'job_entries': len(re.findall(r'^\*\*[^*]+\*\*\s*$', markdown_text, re.MULTILINE)),
            'estimated_words': len(markdown_text.split()),
        }
        
        # Calculate DOCX-specific density score
        density_score = (
            analysis['content_lines'] * 1.2 +
            analysis['h2_headers'] * 2.5 +
            analysis['bullet_points'] * 1.8 +
            analysis['job_entries'] * 3.0 +
            analysis['estimated_words'] * 0.12
        )
        
        analysis['density_score'] = int(density_score)
        return analysis
    
    def calculate_docx_scaling(self, analysis: Dict[str, int]) -> Dict[str, Any]:
        """
        Calculate DOCX font scaling based on content density
        """
        content_lines = analysis['content_lines']
        
        print(f"📊 DOCX Analysis: {content_lines} content lines, {analysis['estimated_words']} words")
        
        # Define DOCX scaling tiers
        if content_lines <= 25:
            # Short resume - comfortable fonts
            scale = {
                'factor': 1.0, 'class': 'COMFORTABLE',
                'header': 18, 'section': 12, 'body': 10, 'small': 9,
                'line_spacing': 1.15, 'spacing_after': 6, 'margins': 0.8
            }
        elif content_lines <= 40:
            # Medium resume - slightly reduced
            scale = {
                'factor': 0.9, 'class': 'BALANCED',
                'header': 16, 'section': 11, 'body': 9, 'small': 8,
                'line_spacing': 1.1, 'spacing_after': 4, 'margins': 0.7
            }
        elif content_lines <= 60:
            # Long resume - compact
            scale = {
                'factor': 0.8, 'class': 'COMPACT',
                'header': 15, 'section': 10, 'body': 8, 'small': 7,
                'line_spacing': 1.05, 'spacing_after': 3, 'margins': 0.6
            }
        elif content_lines <= 80:
            # Very long resume - dense
            scale = {
                'factor': 0.7, 'class': 'DENSE',
                'header': 14, 'section': 9, 'body': 7, 'small': 6,
                'line_spacing': 1.0, 'spacing_after': 2, 'margins': 0.5
            }
        else:
            # Extremely long resume - ultra dense
            scale = {
                'factor': 0.6, 'class': 'ULTRA_DENSE',
                'header': 13, 'section': 9, 'body': 7, 'small': 6,
                'line_spacing': 0.95, 'spacing_after': 1, 'margins': 0.4
            }
        
        print(f"🎯 DOCX Auto-scaling: {scale['class']} ({scale['factor']:.1f}x) - Body: {scale['body']}pt")
        
        return scale
    
    def setup_auto_fit_styles(self, document, scaling: Dict[str, Any]):
        """
        Setup auto-fit styles based on scaling factors
        """
        styles = document.styles
        s = scaling  # Shorthand
        
        # Document defaults with auto-fit sizing
        normal_style = document.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(s['body'])
        normal_font.color.rgb = self.text_color
        
        # Auto-fit paragraph spacing
        normal_format = normal_style.paragraph_format
        normal_format.space_after = Pt(s['spacing_after'])
        normal_format.space_before = Pt(0)
        normal_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_format.line_spacing = s['line_spacing']
        
        # Auto-fit header style
        if 'AutoHeader' not in [st.name for st in styles]:
            header_style = styles.add_style('AutoHeader', WD_STYLE_TYPE.PARAGRAPH)
            header_font = header_style.font
            header_font.name = 'Calibri'
            header_font.size = Pt(s['header'])
            header_font.bold = True
            header_font.color.rgb = self.primary_color
            header_format = header_style.paragraph_format
            header_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_format.space_after = Pt(max(4, s['spacing_after']))
            header_format.space_before = Pt(0)
        
        # Auto-fit contact info
        if 'AutoContact' not in [st.name for st in styles]:
            contact_style = styles.add_style('AutoContact', WD_STYLE_TYPE.PARAGRAPH)
            contact_font = contact_style.font
            contact_font.name = 'Calibri'
            contact_font.size = Pt(s['small'])
            contact_font.color.rgb = self.light_color
            contact_format = contact_style.paragraph_format
            contact_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            contact_format.space_after = Pt(max(6, s['spacing_after'] + 2))
            contact_format.space_before = Pt(1)
        
        # Auto-fit section headers
        if 'AutoSection' not in [st.name for st in styles]:
            section_style = styles.add_style('AutoSection', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Calibri'
            section_font.size = Pt(s['section'])
            section_font.bold = True
            section_font.color.rgb = self.secondary_color
            section_format = section_style.paragraph_format
            section_format.space_after = Pt(max(2, s['spacing_after'] // 2))
            section_format.space_before = Pt(max(6, s['spacing_after'] + 2))
        
        # Auto-fit job title
        if 'AutoJobTitle' not in [st.name for st in styles]:
            job_style = styles.add_style('AutoJobTitle', WD_STYLE_TYPE.PARAGRAPH)
            job_font = job_style.font
            job_font.name = 'Calibri'
            job_font.size = Pt(s['body'])
            job_font.bold = True
            job_font.color.rgb = self.text_color
            job_style.paragraph_format.space_after = Pt(1)
            job_style.paragraph_format.space_before = Pt(max(2, s['spacing_after'] // 2))
        
        # Auto-fit company info
        if 'AutoCompany' not in [st.name for st in styles]:
            company_style = styles.add_style('AutoCompany', WD_STYLE_TYPE.PARAGRAPH)
            company_font = company_style.font
            company_font.name = 'Calibri'
            company_font.size = Pt(s['small'])
            company_font.italic = True
            company_font.color.rgb = self.light_color
            company_style.paragraph_format.space_after = Pt(max(2, s['spacing_after'] // 2))
            company_style.paragraph_format.space_before = Pt(0)
        
        # Auto-fit list style
        if 'AutoList' not in [st.name for st in styles]:
            list_style = styles.add_style('AutoList', WD_STYLE_TYPE.PARAGRAPH)
            list_font = list_style.font
            list_font.name = 'Calibri'
            list_font.size = Pt(s['body'])
            list_font.color.rgb = self.text_color
            list_format = list_style.paragraph_format
            list_format.space_after = Pt(max(1, s['spacing_after'] // 3))
            list_format.space_before = Pt(0)
            list_format.left_indent = Inches(0.15)
            list_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            list_format.line_spacing = s['line_spacing']
    
    def add_auto_section_border(self, paragraph):
        """Add subtle border to auto-fit section headers"""
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pBdr.set(qn('w:bottom'), 'single')
        pBdr.set(qn('w:sz'), '3')
        pBdr.set(qn('w:space'), '1')
        pBdr.set(qn('w:color'), '2c5aa0')
        pPr.append(pBdr)
    
    def parse_auto_fit_text(self, text: str, paragraph):
        """Parse text with auto-fit rich formatting"""
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
        
        for part in parts:
            if not part:
                continue
                
            run = paragraph.add_run()
            
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
                run.text = part[1:-1]
                run.italic = True
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                run.text = part[1:-1]
                run.font.name = 'Consolas'
                run.font.size = Pt(7)  # Small code font
            elif '[' in part and '](' in part:
                link_match = re.match(r'\[([^\]]+)\]\([^\)]+\)', part)
                if link_match:
                    run.text = link_match.group(1)
                    run.font.color.rgb = RGBColor(49, 130, 206)
                else:
                    run.text = part
            else:
                run.text = part
    
    def fix_auto_fit_markdown(self, markdown_text: str) -> str:
        """Fix markdown for auto-fit processing"""
        text = markdown_text.strip()
        
        # Aggressive blank line removal for auto-fit
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'\n\n+', '\n\n', text)
        
        # Clean whitespace
        text = re.sub(r'[ \t]+$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^[ \t]+', '', text, flags=re.MULTILINE)
        
        return text.strip()
    
    def generate_auto_fit_docx_base64(self, markdown_text: str) -> str:
        """
        Generate auto-fit DOCX with dynamic sizing for single page
        """
        try:
            print("📝 Starting AUTO-FIT DOCX generation...")
            
            # Analyze content and calculate scaling
            analysis = self.analyze_docx_content(markdown_text)
            scaling = self.calculate_docx_scaling(analysis)
            
            # Fix markdown
            clean_markdown = self.fix_auto_fit_markdown(markdown_text)
            
            # Create document with auto-fit margins
            document = Document()
            
            # Auto-fit margins based on scaling
            for section in document.sections:
                margin = scaling['margins']
                section.top_margin = Inches(margin)
                section.bottom_margin = Inches(margin) 
                section.left_margin = Inches(margin)
                section.right_margin = Inches(margin)
            
            # Setup auto-fit styles
            self.setup_auto_fit_styles(document, scaling)
            
            # Parse content with auto-fit formatting
            lines = clean_markdown.split('\n')
            skip_next_empty = False
            
            for i, line in enumerate(lines):
                line = line.strip()
                
                # Skip empty lines more aggressively for auto-fit
                if not line:
                    if skip_next_empty or scaling['class'] in ['DENSE', 'ULTRA_DENSE']:
                        skip_next_empty = False
                        continue
                    continue
                
                # Main header (name)
                if line.startswith('# '):
                    header_text = line[2:].strip()
                    p = document.add_paragraph(header_text, style='AutoHeader')
                    skip_next_empty = True
                    continue
                
                # Section headers
                elif line.startswith('## '):
                    section_title = line[3:].strip()
                    p = document.add_paragraph(section_title, style='AutoSection')
                    self.add_auto_section_border(p)
                    skip_next_empty = True
                    continue
                
                # Contact info
                elif any(indicator in line.lower() for indicator in ['@', 'linkedin', 'github', 'phone', '(', ')', '+', 'http']):
                    if not line.startswith(('*', '-', '1', '2', '3', '4', '5')):
                        contact_p = document.add_paragraph()
                        contact_p.style = 'AutoContact'
                        self.parse_auto_fit_text(line, contact_p)
                        continue
                
                # Job titles
                elif line.startswith('**') and line.endswith('**'):
                    job_title = line[2:-2].strip()
                    p = document.add_paragraph(job_title, style='AutoJobTitle')
                    continue
                
                # Company info
                elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
                    company_info = line[1:-1].strip()
                    p = document.add_paragraph(company_info, style='AutoCompany')
                    continue
                
                # Bullet points
                elif line.startswith(('- ', '• ', '* ')):
                    bullet_text = re.sub(r'^[-•*]\s+', '', line)
                    p = document.add_paragraph(style='AutoList')
                    p.add_run('• ')
                    self.parse_auto_fit_text(bullet_text, p)
                    continue
                
                # Regular paragraphs
                else:
                    p = document.add_paragraph()
                    p.style = 'Normal'
                    # Extra tight spacing for dense layouts
                    if scaling['class'] in ['DENSE', 'ULTRA_DENSE']:
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.space_before = Pt(0)
                    self.parse_auto_fit_text(line, p)
            
            # Save to bytes
            doc_buffer = io.BytesIO()
            document.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Encode to base64
            docx_base64 = base64.b64encode(doc_buffer.read()).decode('utf-8')
            print(f"✅ AUTO-FIT DOCX generated! Size: {len(docx_base64)} characters")
            print(f"📄 Single-page layout with {scaling['class']} formatting")
            
            return docx_base64
            
        except Exception as e:
            print(f"❌ Auto-fit DOCX generation failed: {e}")
            import traceback
            traceback.print_exc()
            raise

# Test the auto-fit DOCX generator
def test_auto_fit_docx():
    """Test auto-fit DOCX with different content lengths"""
    
    test_resumes = {
        "Short": """# Jane Doe
jane@email.com | (555) 123-4567

## Experience
**Software Engineer**
*Tech Co | 2020 - Present*
- Built applications

## Skills
Python, JavaScript""",
        
        "Long": """# Dr. Michael Chen
michael.chen@company.com | (555) 987-6543 | LinkedIn: linkedin.com/in/michael-chen

## Professional Summary
**Senior Engineering Manager** with **15+ years** experience leading *technical teams*. Expert in **distributed systems**, `microservices`, and *cloud architecture*. Led teams of **30+ engineers** across multiple products.

## Experience

**Senior Engineering Manager**
*Google | Mountain View, CA | 2018 - Present*
- Lead engineering teams for **Google Cloud** platform
- **Responsibilities**: Architecture design, team management, product strategy
- **Achievements**: Launched **5 major products**, improved team velocity by **60%**
- **Technologies**: `Kubernetes`, **Go**, *Python*, `gRPC`, **Istio**

**Principal Software Engineer** 
*Netflix | Los Gatos, CA | 2014 - 2018*
- Architected **streaming infrastructure** serving `200M+ users`
- **Performance**: Reduced latency by **45%** through optimization
- **Scalability**: Built systems handling **100K+ requests/second**

**Senior Software Engineer**
*Facebook | Menlo Park, CA | 2010 - 2014*
- Developed **social graph** algorithms and **recommendation systems**
- **Impact**: Increased user engagement by **30%**

## Education

**Master of Science in Computer Science**
*Stanford University | Stanford, CA | 2008 - 2010*
- **Specialization**: Distributed Systems and Machine Learning
- **Research**: Published **3 papers** in top-tier conferences

**Bachelor of Engineering in Computer Science**
*University of California, Berkeley | Berkeley, CA | 2004 - 2008*

## Technical Skills

**Programming Languages**: `Python`, **Java**, `Go`, **C++**, *JavaScript*
**Frameworks**: **Django**, `Flask`, *Spring Boot*, **React**, `Node.js`
**Databases**: `PostgreSQL`, **MongoDB**, *Cassandra*, **Redis**, `MySQL`
**Cloud Platforms**: **AWS**, `Google Cloud`, *Azure*, **Docker**, `Kubernetes`
**Tools**: **Git**, `Jenkins`, *Terraform*, **Prometheus**, `Grafana`"""
    }
    
    print("🧪 Testing AUTO-FIT DOCX with Different Lengths")
    print("=" * 50)
    
    generator = AutoFitDocxGenerator()
    
    for name, content in test_resumes.items():
        print(f"\n📝 Testing {name} Resume:")
        print("-" * 30)
        
        try:
            result = generator.generate_auto_fit_docx_base64(content)
            print(f"✅ Success: {len(result):,} characters")
        except Exception as e:
            print(f"❌ Failed: {e}")
    
    print("\n🎉 AUTO-FIT DOCX TEST COMPLETE!")

if __name__ == "__main__":
    test_auto_fit_docx()